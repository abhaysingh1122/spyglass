"""Word-document builder — built from STRUCTURED analysis data, never raw AI markdown."""
import os
import re
import datetime as dt
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "reports")

NAVY = RGBColor(0x1A, 0x2B, 0x4A)
GOLD = RGBColor(0xB8, 0x8A, 0x2E)
GREY = RGBColor(0x6B, 0x6B, 0x6B)


def _strip_md(text: str) -> str:
    """Remove markdown artifacts if any leak through."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text or "")
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    return text.replace("---", "").strip()


def _label_value(doc, label, value, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(size)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(_strip_md(str(value or "—")))
    r2.font.size = Pt(size)


def _ago(posted_iso) -> str:
    try:
        posted = dt.datetime.fromisoformat(str(posted_iso).replace("Z", "+00:00"))
        hours = int((dt.datetime.now(dt.timezone.utc) - posted).total_seconds() // 3600)
        return "just now" if hours < 1 else (f"{hours}h ago" if hours < 24 else f"{hours//24}d ago")
    except Exception:
        return ""


def build_docx(result: dict, new_posts: list, growth_updates: list,
               competitor_names: dict = None) -> str:
    os.makedirs(OUT_DIR, exist_ok=True)
    today = dt.date.today()
    path = os.path.join(OUT_DIR, f"SpyGlass Intel {today.isoformat()}.docx")
    by_url = {p.get("post_url"): p for p in new_posts}
    names = competitor_names or {}

    doc = Document()

    # ---- Title block ----
    title = doc.add_paragraph()
    r = title.add_run("SPYGLASS")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY
    r2 = title.add_run("  ·  Daily Competitor Intel")
    r2.font.size = Pt(16)
    r2.font.color.rgb = GOLD

    meta = doc.add_paragraph()
    m = meta.add_run(f"{today.strftime('%A, %B %d, %Y')}   |   "
                     f"{len(new_posts)} new post(s)   |   {len(growth_updates)} growth re-check(s)")
    m.font.size = Pt(10)
    m.font.color.rgb = GREY
    meta.paragraph_format.space_after = Pt(16)

    # ---- Executive read ----
    if result.get("overall_pattern"):
        h = doc.add_heading("The Pattern", level=1)
        h.runs[0].font.color.rgb = NAVY
        p = doc.add_paragraph(_strip_md(result["overall_pattern"]))
        p.runs[0].font.size = Pt(11)
        p.runs[0].italic = True
        p.paragraph_format.space_after = Pt(14)

    # ---- Per-post breakdown (structured, never raw AI text) ----
    if result.get("posts"):
        h = doc.add_heading("New Posts — Last 24 Hours", level=1)
        h.runs[0].font.color.rgb = NAVY
        for i, p in enumerate(result["posts"], 1):
            row = by_url.get(p.get("post_url"), {})
            comp = names.get(row.get("competitor_id"), "")

            # Post heading
            ph = doc.add_heading(f"Post {i} — {_strip_md(p.get('one_liner', 'New post'))}", level=2)
            ph.runs[0].font.color.rgb = GOLD

            # Meta line
            bits = [b for b in [comp, _ago(row.get("posted_at")),
                                f"{row.get('likes', 0):,} likes",
                                f"{row.get('comments', 0):,} comments",
                                f"{row.get('shares', 0):,} shares"] if b]
            ml = doc.add_paragraph()
            mr = ml.add_run("  ·  ".join(bits))
            mr.font.size = Pt(9)
            mr.font.color.rgb = GREY
            ml.paragraph_format.space_after = Pt(8)

            _label_value(doc, "AI's Take:", p.get("ai_take"), size=11)
            _label_value(doc, "Steal This:", p.get("steal_this"))
            _label_value(doc, "Hook:", f'"{_strip_md(p.get("hook", ""))}"  ({p.get("hook_type", "?")})')
            _label_value(doc, "Content Type:", p.get("content_type"))
            _label_value(doc, "Strategy:", p.get("strategy"))
            _label_value(doc, "Audience Reaction:", p.get("audience_reaction"))

            # Content excerpt
            if row.get("content"):
                ex = doc.add_paragraph()
                er = ex.add_run("“" + row["content"][:400].strip() + ("…" if len(row["content"]) > 400 else "") + "”")
                er.font.size = Pt(9)
                er.italic = True
                er.font.color.rgb = GREY
                ex.paragraph_format.left_indent = Inches(0.3)

            # Link
            lk = doc.add_paragraph()
            lr = lk.add_run(p.get("post_url", ""))
            lr.font.size = Pt(8)
            lr.font.color.rgb = RGBColor(0x2E, 0x6D, 0xB4)
            lk.paragraph_format.space_after = Pt(14)

    # ---- Growth section ----
    if growth_updates or result.get("growth_lines"):
        h = doc.add_heading("Growth Watch — 7-Day Re-checks", level=1)
        h.runs[0].font.color.rgb = NAVY
        for line in result.get("growth_lines", []) or []:
            b = doc.add_paragraph(style="List Bullet")
            b.add_run(_strip_md(line)).font.size = Pt(10)
        for g in growth_updates:
            prev, cur = g.get("previous", {}), g.get("current", {})
            b = doc.add_paragraph(style="List Bullet")
            br = b.add_run(
                f"{g.get('post_url', '')} — likes {prev.get('likes', 0):,} → {cur.get('likes', 0):,}, "
                f"comments {prev.get('comments', 0):,} → {cur.get('comments', 0):,}, "
                f"shares {prev.get('shares', 0):,} → {cur.get('shares', 0):,}")
            br.font.size = Pt(9)

    # ---- Footer ----
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = f.add_run("— SpyGlass is watching —")
    fr.font.size = Pt(9)
    fr.font.color.rgb = GOLD

    doc.save(path)
    return path


def build_dossier_docx(name: str, result: dict, posts: list,
                       competitor_names: dict = None) -> str:
    """The Content-Spy dossier document — full playbook decode."""
    os.makedirs(OUT_DIR, exist_ok=True)
    today = dt.date.today()
    path = os.path.join(OUT_DIR, f"SpyGlass Dossier - {name.title()} {today.isoformat()}.docx")

    doc = Document()
    t = doc.add_paragraph()
    r = t.add_run("SPYGLASS")
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY
    r2 = t.add_run(f"  ·  Content Dossier: {name.title()}")
    r2.font.size = Pt(16); r2.font.color.rgb = GOLD

    meta = doc.add_paragraph()
    m = meta.add_run(f"{today.strftime('%B %d, %Y')}   |   compiled from {len(posts)} posts")
    m.font.size = Pt(10); m.font.color.rgb = GREY
    meta.paragraph_format.space_after = Pt(14)

    h = doc.add_heading("Verdict", level=1); h.runs[0].font.color.rgb = NAVY
    v = doc.add_paragraph(_strip_md(result.get("verdict", "—")))
    v.runs[0].italic = True; v.runs[0].font.size = Pt(12)

    if result.get("hook_matrix"):
        h = doc.add_heading("Hook Matrix", level=1); h.runs[0].font.color.rgb = NAVY
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, col in enumerate(["Hook Type", "Times Used", "Avg Engagement", "Verdict"]):
            hdr[i].paragraphs[0].add_run(col).bold = True
        for hm in result["hook_matrix"]:
            row = table.add_row().cells
            row[0].text = str(hm.get("hook_type", ""))
            row[1].text = str(hm.get("times_used", ""))
            row[2].text = f"{hm.get('avg_engagement', 0):,}"
            row[3].text = str(hm.get("verdict", ""))
        doc.add_paragraph()

    if result.get("content_mix"):
        h = doc.add_heading("Content Mix", level=1); h.runs[0].font.color.rgb = NAVY
        for c in result["content_mix"]:
            b = doc.add_paragraph(style="List Bullet")
            br = b.add_run(f"{c.get('content_type', '?')} — {c.get('share_pct', 0)}%  ")
            br.bold = True
            b.add_run(str(c.get("note", "")))

    if result.get("cadence"):
        _label_value(doc, "Cadence:", result["cadence"], size=11)

    tp = result.get("top_post") or {}
    if tp:
        h = doc.add_heading("Their Best Weapon", level=1); h.runs[0].font.color.rgb = NAVY
        _label_value(doc, "Post:", f"{tp.get('one_liner', '')} ({tp.get('engagement', '')})", size=11)
        _label_value(doc, "Why it won:", tp.get("why_it_won"), size=11)

    if result.get("weaknesses"):
        h = doc.add_heading("Exploitable Gaps", level=1); h.runs[0].font.color.rgb = NAVY
        for w in result["weaknesses"]:
            doc.add_paragraph(_strip_md(w), style="List Bullet")

    if result.get("playbook"):
        h = doc.add_heading("Steal-This Playbook", level=1); h.runs[0].font.color.rgb = NAVY
        for p in result["playbook"]:
            doc.add_paragraph(_strip_md(p), style="List Number")

    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = f.add_run("— SpyGlass is watching —")
    fr.font.size = Pt(9); fr.font.color.rgb = GOLD

    doc.save(path)
    return path
